from __future__ import annotations

from pathlib import Path
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

BODY_FONT = "Aptos"
CJK_FONT = "PingFang SC"
TITLE_FONT = "Aptos Display"
BODY_SIZE = Pt(10.5)
TABLE_SIZE = Pt(9)
TITLE_SIZE = Pt(18)
HEADING1_SIZE = Pt(13)
HEADING2_SIZE = Pt(11.5)


def contains_cjk(text: str) -> bool:
    return bool(re.search(r"[\u4e00-\u9fff]", text))


def english_safe_text(text: str, fallback: str) -> str:
    candidate = str(text or "").strip()
    if not candidate:
        return fallback
    if contains_cjk(candidate):
        return fallback
    return candidate


def build_english_report_fallback(payload: dict[str, Any]) -> dict[str, Any]:
    ticker = str(payload.get("ticker") or "UNKNOWN").strip() or "UNKNOWN"
    company_name = str(payload.get("company_name") or "").strip()
    safe_company_name = ticker if contains_cjk(company_name) else (company_name or ticker)
    verdict = english_safe_text(str(payload.get("verdict") or ""), "watchlist")
    confidence = english_safe_text(str(payload.get("confidence") or ""), "medium")
    evidence_count = len(payload.get("evidence") or []) if isinstance(payload.get("evidence"), list) else 0
    target_snapshot = _english_target_snapshot(payload.get("target_prices") or {})
    return {
        "company_name": safe_company_name,
        "ticker": ticker,
        "exchange": english_safe_text(str(payload.get("exchange") or ""), "n/a"),
        "market": english_safe_text(str(payload.get("market") or ""), "n/a"),
        "quick_take": english_safe_text(
            str(payload.get("quick_take") or ""),
            f"Chinese report is the primary memo. English fallback summary: verdict={verdict}, confidence={confidence}, evidence sources={evidence_count}, target snapshot={target_snapshot}.",
        ),
        "verdict": verdict,
        "confidence": confidence,
        "recent_developments": english_safe_text(str(payload.get("recent_developments") or ""), "Recent developments remain under review; use the Chinese memo and evidence table for the freshest source trail."),
        "market_map": english_safe_text(str(payload.get("market_map") or ""), "Market context remains under review; rely on the Chinese memo and evidence table for the detailed source trail."),
        "business_summary": english_safe_text(str(payload.get("business_summary") or ""), f"{safe_company_name} is being treated as a watchlist research candidate until business mix, customer quality, and order durability are verified."),
        "china_story": english_safe_text(str(payload.get("china_story") or ""), "The strategic narrative should be tested against concrete orders, customer concentration, and margin evidence before being treated as conviction."),
        "sentiment_simulation": english_safe_text(str(payload.get("sentiment_simulation") or ""), "Sentiment remains two-sided: the bull case trades optionality, while the bear case focuses on proof gaps and valuation risk."),
        "peer_comparison": english_safe_text(str(payload.get("peer_comparison") or ""), "Peer comparison requires a cleaner valuation anchor and higher-quality comparable-company evidence."),
        "committee_takeaways": english_safe_text(str(payload.get("committee_takeaways") or ""), "The council fallback keeps the name on watchlist rather than high conviction until stronger primary evidence arrives."),
        "scenario_outlook": english_safe_text(str(payload.get("scenario_outlook") or ""), "Base case remains watchlist; bull case needs verified order acceleration; bear case is narrative overshoot without fundamental confirmation."),
        "debate_notes": english_safe_text(str(payload.get("debate_notes") or ""), "Red-team focus: customer concentration, evidence quality, order durability, and whether the valuation anchor is robust."),
        "valuation_view": english_safe_text(str(payload.get("valuation_view") or ""), f"Fallback valuation view: {target_snapshot}. Treat these as scenario targets, not investment advice."),
        "macro_context": english_safe_text(str(payload.get("macro_context") or ""), "Macro context translation unavailable."),
        "flow_signal": english_safe_text(str(payload.get("flow_signal") or ""), "Flow signal translation unavailable."),
        "technical_view": english_safe_text(str(payload.get("technical_view") or ""), "Technical view translation unavailable."),
        "factor_exposure": payload.get("factor_exposure") or {},
        "catalyst_calendar": payload.get("catalyst_calendar") or [],
        "bull_case": _english_safe_list(payload.get("bull_case"), "Bull case: upside requires verified order quality, customer expansion, and durable margin or mix improvement."),
        "bear_case": _english_safe_list(payload.get("bear_case"), "Bear case: the narrative may outrun evidence if customer concentration, order durability, or valuation anchors remain weak."),
        "catalysts": _english_safe_list(payload.get("catalysts"), "Catalysts: new orders, customer proof, margin evidence, and sector capital-expenditure recovery."),
        "risks": _english_safe_list(payload.get("risks"), "Risks: weak evidence quality, customer concentration, demand cyclicality, and valuation compression."),
        "next_questions": _english_safe_list(payload.get("next_questions"), "Next question: what primary filings or company disclosures can verify the key investment debate?"),
        "evidence": _english_safe_evidence(payload.get("evidence")),
        "target_prices": _english_safe_targets(payload.get("target_prices")),
    }


def build_screening_doc_payload(*, theme: str, market: str, stage_one_candidates: list[dict[str, Any]], finalists: list[dict[str, Any]]) -> dict[str, Any]:
    rejected = [
        item
        for item in stage_one_candidates
        if _candidate_identity(item) not in {_candidate_identity(finalist) for finalist in finalists}
    ]
    return {
        "theme": theme,
        "market": market,
        "generated_at": "",
        "stage_one_candidates": [
            {
                "company_name": str(item.get("company_name") or "").strip(),
                "ticker": str(item.get("ticker") or "").strip(),
                "screen_score": item.get("screen_score"),
                "why_now": str(item.get("why_now") or item.get("rationale") or "").strip(),
                "vertical_summary": str(item.get("vertical_summary") or "").strip(),
                "horizontal_summary": str(item.get("horizontal_summary") or "").strip(),
                "why_not_now": str(item.get("why_not_now") or "").strip(),
            }
            for item in stage_one_candidates
        ],
        "finalists": [
            {
                "company_name": str(item.get("company_name") or "").strip(),
                "ticker": str(item.get("ticker") or "").strip(),
                "screen_score": item.get("screen_score"),
                "recommendation_rank": str(item.get("recommendation_rank") or "").strip(),
                "research_verdict": str((item.get("payload") or {}).get("verdict") or "").strip(),
                "confidence": str((item.get("payload") or {}).get("confidence") or "").strip(),
                "why_now": str(item.get("stage_two_note") or item.get("rationale") or "").strip(),
                "why_not_now": str(item.get("why_not_now") or "").strip(),
                "quick_take": str((item.get("payload") or {}).get("quick_take") or "").strip(),
                "vertical_summary": str(item.get("vertical_summary") or "").strip(),
                "horizontal_summary": str(item.get("horizontal_summary") or "").strip(),
                "bull_bear_focus": str(summarize_bull_bear_for_doc(item.get("payload") or {})).strip(),
                "short_term_target": _target_snapshot((item.get("payload") or {}).get("target_prices") or {}, "short_term"),
                "document_path": str(item.get("primary_document_path") or item.get("zh_docx_path") or item.get("markdown_path") or "").strip(),
            }
            for item in finalists
        ],
        "rejected": [
            {
                "label": str(item.get("ticker") or item.get("company_name") or "").strip(),
                "reason": str(item.get("exclusion_reason") or item.get("why_not_now") or item.get("rationale") or "research upside was weaker").strip(),
            }
            for item in rejected[:8]
        ],
    }


def build_english_screening_fallback(payload: dict[str, Any]) -> dict[str, Any]:
    return {
        "theme": english_safe_text(str(payload.get("theme") or ""), "theme"),
        "market": english_safe_text(str(payload.get("market") or ""), "market"),
        "generated_at": str(payload.get("generated_at") or ""),
        "stage_one_candidates": [
            {
                "company_name": _safe_company_label(item),
                "ticker": str(item.get("ticker") or "").strip(),
                "screen_score": item.get("screen_score"),
                "why_now": english_safe_text(str(item.get("why_now") or ""), "Why-now translation was unavailable for this candidate."),
                "vertical_summary": english_safe_text(str(item.get("vertical_summary") or ""), "Vertical diligence translation was unavailable."),
                "horizontal_summary": english_safe_text(str(item.get("horizontal_summary") or ""), "Horizontal diligence translation was unavailable."),
                "why_not_now": english_safe_text(str(item.get("why_not_now") or ""), "A cleaner English rejection note was unavailable in this run."),
            }
            for item in payload.get("stage_one_candidates") or []
        ],
        "finalists": [
            {
                "company_name": _safe_company_label(item),
                "ticker": str(item.get("ticker") or "").strip(),
                "screen_score": item.get("screen_score"),
                "recommendation_rank": english_safe_text(str(item.get("recommendation_rank") or ""), "A"),
                "research_verdict": english_safe_text(str(item.get("research_verdict") or ""), "watchlist"),
                "confidence": english_safe_text(str(item.get("confidence") or ""), "medium"),
                "why_now": english_safe_text(str(item.get("why_now") or ""), "Why-now translation was unavailable for this finalist."),
                "why_not_now": english_safe_text(str(item.get("why_not_now") or ""), "A clearer English downgrade note was unavailable in this run."),
                "quick_take": english_safe_text(str(item.get("quick_take") or ""), "Quick take translation was unavailable in this run."),
                "vertical_summary": english_safe_text(str(item.get("vertical_summary") or ""), "Vertical diligence translation was unavailable."),
                "horizontal_summary": english_safe_text(str(item.get("horizontal_summary") or ""), "Horizontal diligence translation was unavailable."),
                "bull_bear_focus": english_safe_text(str(item.get("bull_bear_focus") or ""), "Bull/Bear focus translation was unavailable."),
                "short_term_target": english_safe_text(str(item.get("short_term_target") or ""), "n/a"),
                "document_path": str(item.get("document_path") or "").strip(),
            }
            for item in payload.get("finalists") or []
        ],
        "rejected": [
            {
                "label": english_safe_text(str(item.get("label") or ""), "candidate"),
                "reason": english_safe_text(str(item.get("reason") or ""), "English explanation was unavailable for this excluded name."),
            }
            for item in payload.get("rejected") or []
        ],
    }


def write_report_docx(path: Path, *, payload: dict[str, Any], language: str) -> None:
    doc = _new_document()
    _append_report_section(doc, payload=payload, language=language)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def write_screening_docx(path: Path, *, payload: dict[str, Any], language: str) -> None:
    doc = _new_document()
    _append_screening_section(doc, payload=payload, language=language)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def write_watchlist_digest_docx(path: Path, *, artifacts: list[dict[str, str]], language: str) -> None:
    doc = _new_document()
    _append_watchlist_digest_section(doc, artifacts=artifacts, language=language)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def write_bilingual_report_docx(path: Path, *, zh_payload: dict[str, Any], en_payload: dict[str, Any]) -> None:
    doc = _new_document()
    _add_title(
        doc,
        f"{zh_payload.get('company_name') or zh_payload.get('ticker') or 'Stock'} Research Desk Delivery",
        subtitle="Chinese section first. English section follows on a separate page.",
    )
    _format_paragraph(doc.add_heading("中文交付", level=1), size=HEADING1_SIZE, bold=True)
    _append_report_section(doc, payload=zh_payload, language="zh", include_title=False)
    doc.add_page_break()
    _format_paragraph(doc.add_heading("English Delivery", level=1), size=HEADING1_SIZE, bold=True)
    _append_report_section(doc, payload=en_payload, language="en", include_title=False)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def write_bilingual_screening_docx(path: Path, *, zh_payload: dict[str, Any], en_payload: dict[str, Any]) -> None:
    doc = _new_document()
    _add_title(
        doc,
        f"{zh_payload.get('theme', '')} Screening Delivery",
        subtitle="Chinese section first. English section follows on a separate page.",
    )
    _format_paragraph(doc.add_heading("中文交付", level=1), size=HEADING1_SIZE, bold=True)
    _append_screening_section(doc, payload=zh_payload, language="zh", include_title=False)
    doc.add_page_break()
    _format_paragraph(doc.add_heading("English Delivery", level=1), size=HEADING1_SIZE, bold=True)
    _append_screening_section(doc, payload=en_payload, language="en", include_title=False)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def write_bilingual_watchlist_digest_docx(path: Path, *, artifacts: list[dict[str, str]]) -> None:
    doc = _new_document()
    _add_title(
        doc,
        "Watchlist Delivery",
        subtitle="Chinese section first. English section follows on a separate page.",
    )
    _format_paragraph(doc.add_heading("中文交付", level=1), size=HEADING1_SIZE, bold=True)
    _append_watchlist_digest_section(doc, artifacts=artifacts, language="zh", include_title=False)
    doc.add_page_break()
    _format_paragraph(doc.add_heading("English Delivery", level=1), size=HEADING1_SIZE, bold=True)
    _append_watchlist_digest_section(doc, artifacts=artifacts, language="en", include_title=False)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def summarize_bull_bear_for_doc(payload: dict[str, Any]) -> str:
    bull_case = list(payload.get("bull_case") or [])
    bear_case = list(payload.get("bear_case") or [])
    bull = bull_case[0] if bull_case else "bull case still needs verification"
    bear = bear_case[0] if bear_case else "bear case still needs verification"
    return f"Bull: {bull} | Bear: {bear}"


def _target_snapshot(targets: dict[str, Any], key: str) -> str:
    item = targets.get(key) or {}
    price = str(item.get("price") or "").strip() or "n/a"
    horizon = str(item.get("horizon") or "").strip() or "n/a"
    return f"{price} | {horizon}"


def _english_safe_list(value: Any, fallback: str) -> list[str]:
    if not isinstance(value, list):
        return [fallback]
    cleaned = [english_safe_text(str(item), fallback) for item in value if str(item).strip()]
    return cleaned or [fallback]


def _english_safe_evidence(value: Any) -> list[dict[str, str]]:
    if not isinstance(value, list):
        return []
    normalized: list[dict[str, str]] = []
    for item in value:
        if not isinstance(item, dict):
            continue
        normalized.append(
            {
                "title": english_safe_text(str(item.get("title") or ""), "Evidence source"),
                "url": str(item.get("url") or "").strip(),
                "claim": english_safe_text(str(item.get("claim") or ""), "Claim translation unavailable in this run."),
                "stance": english_safe_text(str(item.get("stance") or ""), "neutral"),
            }
        )
    return normalized


def _english_safe_targets(value: Any) -> dict[str, dict[str, str]]:
    if not isinstance(value, dict):
        return {}
    normalized: dict[str, dict[str, str]] = {}
    for key in ("short_term", "medium_term", "long_term"):
        item = value.get(key) or {}
        if not isinstance(item, dict):
            continue
        normalized[key] = {
            "price": str(item.get("price") or "").strip(),
            "horizon": _english_horizon(str(item.get("horizon") or "").strip()),
            "thesis": english_safe_text(str(item.get("thesis") or ""), _fallback_target_thesis(key)),
        }
    return normalized


def _english_horizon(value: str) -> str:
    text = value.strip()
    replacements = {
        "个月": "months",
        "月": "months",
        "年": "years",
        "周": "weeks",
        "日": "days",
    }
    for source, target in replacements.items():
        text = text.replace(source, target)
    text = re.sub(r"(\d)\s*(months|years|weeks|days)", r"\1 \2", text)
    return english_safe_text(text, "n/a")


def _fallback_target_thesis(key: str) -> str:
    return {
        "short_term": "Near-term target tied to order, customer, and sentiment validation.",
        "medium_term": "Medium-term target tied to business-mix migration and valuation-anchor repair.",
        "long_term": "Long-term target tied to whether the company earns a higher-quality equipment-asset multiple.",
    }.get(key, "Scenario target thesis requires further evidence.")


def _english_target_snapshot(value: Any) -> str:
    targets = _english_safe_targets(value)
    parts: list[str] = []
    for key, label in (("short_term", "short"), ("medium_term", "medium"), ("long_term", "long")):
        item = targets.get(key) or {}
        price = str(item.get("price") or "n/a").strip()
        horizon = str(item.get("horizon") or "n/a").strip()
        parts.append(f"{label}: {price} ({horizon})")
    return "; ".join(parts) if parts else "n/a"


def _safe_company_label(item: dict[str, Any]) -> str:
    company_name = str(item.get("company_name") or "").strip()
    ticker = str(item.get("ticker") or "").strip()
    return ticker if contains_cjk(company_name) else (company_name or ticker or "candidate")


def _candidate_identity(item: dict[str, Any]) -> str:
    return str(item.get("ticker") or item.get("company_name") or "").strip()


def _new_document() -> Document:
    doc = Document()
    _configure_document_styles(doc)
    doc.sections[0].top_margin = Pt(48)
    doc.sections[0].bottom_margin = Pt(48)
    doc.sections[0].left_margin = Pt(54)
    doc.sections[0].right_margin = Pt(54)
    return doc


def _configure_document_styles(doc: Document) -> None:
    _set_style_font(doc.styles["Normal"], size=BODY_SIZE)
    _set_style_font(doc.styles["Heading 1"], size=HEADING1_SIZE, bold=True)
    _set_style_font(doc.styles["Heading 2"], size=HEADING2_SIZE, bold=True)
    _set_style_font(doc.styles["List Bullet"], size=BODY_SIZE)


def _set_style_font(style: Any, *, size: Any, bold: bool | None = None) -> None:
    style.font.name = BODY_FONT
    style.font.size = size
    if bold is not None:
        style.font.bold = bold
    style.element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    style.element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    style.element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)


def _set_run_font(run: Any, *, size: Any | None = None, bold: bool | None = None, italic: bool | None = None, title: bool = False) -> None:
    font_name = TITLE_FONT if title else BODY_FONT
    run.font.name = font_name
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)


def _format_paragraph(paragraph: Any, *, size: Any = BODY_SIZE, bold: bool | None = None) -> None:
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.08
    for run in paragraph.runs:
        _set_run_font(run, size=size, bold=bold)


def _format_table(table: Any) -> None:
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run, size=TABLE_SIZE, bold=(True if row_index == 0 else None))
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.02


def _add_title(doc: Document, title: str, *, subtitle: str = "") -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(title)
    _set_run_font(run, size=TITLE_SIZE, bold=True, title=True)
    if subtitle:
        sub = doc.add_paragraph(subtitle)
        sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_run_font(sub.runs[0], size=Pt(9.5), italic=True)


def _add_metadata_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value or "n/a"
    _format_table(table)


def _add_heading_and_paragraph(doc: Document, heading: str, body: str) -> None:
    _format_paragraph(doc.add_heading(heading, level=1), size=HEADING1_SIZE, bold=True)
    cleaned = str(body or "").strip() or "-"
    for chunk in [part.strip() for part in cleaned.split("\n\n") if part.strip()]:
        _format_paragraph(doc.add_paragraph(chunk), size=BODY_SIZE)


def _add_bullet_section(doc: Document, heading: str, items: list[str], empty_text: str) -> None:
    _format_paragraph(doc.add_heading(heading, level=1), size=HEADING1_SIZE, bold=True)
    if not items:
        _format_paragraph(doc.add_paragraph(empty_text), size=BODY_SIZE)
        return
    for item in items:
        _format_paragraph(doc.add_paragraph(str(item), style="List Bullet"), size=BODY_SIZE)


def _add_target_price_section(doc: Document, labels: dict[str, str], target_prices: dict[str, dict[str, str]]) -> None:
    _format_paragraph(doc.add_heading(labels["target_prices"], level=1), size=HEADING1_SIZE, bold=True)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = labels["horizon"]
    headers[1].text = labels["price"]
    headers[2].text = labels["thesis"]
    for key, title in (
        ("short_term", labels["short_term"]),
        ("medium_term", labels["medium_term"]),
        ("long_term", labels["long_term"]),
    ):
        item = target_prices.get(key) or {}
        row = table.add_row().cells
        row[0].text = f"{title} | {item.get('horizon') or 'n/a'}"
        row[1].text = str(item.get("price") or "n/a")
        row[2].text = str(item.get("thesis") or labels["translation_unavailable"])
    _format_table(table)


def _add_evidence_section(doc: Document, labels: dict[str, str], evidence: list[dict[str, str]]) -> None:
    _format_paragraph(doc.add_heading(labels["evidence"], level=1), size=HEADING1_SIZE, bold=True)
    if not evidence:
        _format_paragraph(doc.add_paragraph(labels["no_items"]), size=BODY_SIZE)
        return
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = labels["source"]
    headers[1].text = labels["claim"]
    headers[2].text = labels["stance"]
    for item in evidence:
        row = table.add_row().cells
        row[0].text = f"{item.get('title') or 'Source'}\n{item.get('url') or ''}"
        row[1].text = str(item.get("claim") or "")
        row[2].text = str(item.get("stance") or "")
    _format_table(table)


def _add_factor_exposure_section(doc: Document, labels: dict[str, str], factor_data: dict[str, Any]) -> None:
    """Add a factor exposure table to the document."""
    _format_paragraph(doc.add_heading(labels["factor_exposure"], level=1))
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = labels["factor_name"]
    headers[1].text = labels["factor_rating"]
    factor_keys = ("value", "momentum", "quality", "size", "volatility")
    for key in factor_keys:
        row = table.add_row().cells
        row[0].text = labels.get(f"factor_{key}", key)
        rating = str((factor_data or {}).get(key) or "n/a")
        row[1].text = rating
    _format_table(table)


def _add_catalyst_calendar_section(doc: Document, labels: dict[str, str], calendar: list[dict[str, str]]) -> None:
    """Add a catalyst calendar table to the document."""
    _format_paragraph(doc.add_heading(labels["catalyst_calendar"], level=1))
    if not calendar:
        p = doc.add_paragraph(labels.get("no_items", "No items."))
        _format_paragraph(p, size=BODY_SIZE)
        return
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = labels["catalyst_event"]
    headers[1].text = labels["catalyst_date"]
    headers[2].text = labels["catalyst_impact"]
    headers[3].text = labels["catalyst_direction"]
    for item in calendar[:10]:
        row = table.add_row().cells
        row[0].text = str(item.get("event", ""))
        row[1].text = str(item.get("date", ""))
        row[2].text = str(item.get("impact", ""))
        row[3].text = str(item.get("direction", ""))
    _format_table(table)


def _append_report_section(doc: Document, *, payload: dict[str, Any], language: str, include_title: bool = True) -> None:
    labels = _report_labels(language)
    if include_title:
        title = f"{payload.get('company_name') or payload.get('ticker')} {labels['memo_title']}"
        _add_title(doc, title, subtitle=f"{labels['ticker']}: {payload.get('ticker', '')} | {labels['market']}: {payload.get('market', '')}")
    _add_metadata_table(
        doc,
        [
            (labels["verdict"], str(payload.get("verdict") or "")),
            (labels["confidence"], str(payload.get("confidence") or "")),
            (labels["exchange"], str(payload.get("exchange") or "")),
            (labels["model"], str(payload.get("model") or "")),
        ],
    )
    _add_heading_and_paragraph(doc, labels["quick_take"], str(payload.get("quick_take") or ""))
    _add_heading_and_paragraph(doc, labels["recent_developments"], str(payload.get("recent_developments") or ""))
    _add_heading_and_paragraph(doc, labels["business_summary"], str(payload.get("business_summary") or ""))
    _add_heading_and_paragraph(doc, labels["market_map"], str(payload.get("market_map") or ""))
    _add_heading_and_paragraph(doc, labels["china_story"], str(payload.get("china_story") or ""))
    _add_heading_and_paragraph(doc, labels["macro_context"], str(payload.get("macro_context") or ""))
    _add_heading_and_paragraph(doc, labels["flow_signal"], str(payload.get("flow_signal") or ""))
    _add_heading_and_paragraph(doc, labels["sentiment_simulation"], str(payload.get("sentiment_simulation") or ""))
    _add_heading_and_paragraph(doc, labels["peer_comparison"], str(payload.get("peer_comparison") or ""))
    _add_heading_and_paragraph(doc, labels["committee_takeaways"], str(payload.get("committee_takeaways") or ""))
    _add_heading_and_paragraph(doc, labels["scenario_outlook"], str(payload.get("scenario_outlook") or ""))
    _add_bullet_section(doc, labels["bull_case"], list(payload.get("bull_case") or []), labels["no_items"])
    _add_bullet_section(doc, labels["bear_case"], list(payload.get("bear_case") or []), labels["no_items"])
    _add_bullet_section(doc, labels["catalysts"], list(payload.get("catalysts") or []), labels["no_items"])
    _add_catalyst_calendar_section(doc, labels, list(payload.get("catalyst_calendar") or []))
    _add_bullet_section(doc, labels["risks"], list(payload.get("risks") or []), labels["no_items"])
    _add_heading_and_paragraph(doc, labels["valuation_view"], str(payload.get("valuation_view") or ""))
    _add_heading_and_paragraph(doc, labels["technical_view"], str(payload.get("technical_view") or ""))
    _add_factor_exposure_section(doc, labels, payload.get("factor_exposure") or {})
    _add_target_price_section(doc, labels, payload.get("target_prices") or {})
    _add_heading_and_paragraph(doc, labels["debate_notes"], str(payload.get("debate_notes") or ""))
    _add_evidence_section(doc, labels, list(payload.get("evidence") or []))
    _add_bullet_section(doc, labels["next_questions"], list(payload.get("next_questions") or []), labels["no_items"])


def _append_screening_section(doc: Document, *, payload: dict[str, Any], language: str, include_title: bool = True) -> None:
    labels = _screening_labels(language)
    if include_title:
        _add_title(doc, f"{payload.get('theme', '')} {labels['title']}", subtitle=f"{labels['market']}: {payload.get('market', '')}")
    finalists = list(payload.get("finalists") or [])
    stage_one = list(payload.get("stage_one_candidates") or [])
    _add_metadata_table(
        doc,
        [
            (labels["second_screen_pool"], str(len(stage_one))),
            (labels["final_recommendations"], str(len(finalists))),
        ],
    )
    _format_paragraph(doc.add_heading(labels["final_recommendations"], level=1), size=HEADING1_SIZE, bold=True)
    if not finalists:
        _format_paragraph(doc.add_paragraph(labels["no_items"]), size=BODY_SIZE)
    for item in finalists:
        _format_paragraph(doc.add_heading(f"{item.get('company_name') or item.get('ticker', '')} {item.get('ticker', '')}".strip(), level=2), size=HEADING2_SIZE, bold=True)
        for line in [
            f"{labels['recommendation_rank']}: {item.get('recommendation_rank') or 'A'}",
            f"{labels['screen_score']}: {item.get('screen_score') or 'n/a'}",
            f"{labels['research_verdict']}: {item.get('research_verdict') or 'watchlist'}",
            f"{labels['confidence']}: {item.get('confidence') or 'medium'}",
            f"{labels['why_now']}: {item.get('why_now') or labels['translation_unavailable']}",
            f"{labels['why_not_now']}: {item.get('why_not_now') or labels['translation_unavailable']}",
            f"{labels['quick_take']}: {item.get('quick_take') or labels['translation_unavailable']}",
            f"{labels['vertical_summary']}: {item.get('vertical_summary') or labels['translation_unavailable']}",
            f"{labels['horizontal_summary']}: {item.get('horizontal_summary') or labels['translation_unavailable']}",
            f"{labels['short_term_target']}: {item.get('short_term_target') or 'n/a'}",
            f"{labels['bull_bear_focus']}: {item.get('bull_bear_focus') or labels['translation_unavailable']}",
            f"{labels['document_path']}: {item.get('document_path') or 'n/a'}",
        ]:
            _format_paragraph(doc.add_paragraph(line, style="List Bullet"), size=BODY_SIZE)
    _format_paragraph(doc.add_heading(labels["downgraded_names"], level=1), size=HEADING1_SIZE, bold=True)
    rejected = list(payload.get("rejected") or [])
    if not rejected:
        _format_paragraph(doc.add_paragraph(labels["no_items"]), size=BODY_SIZE)
    for item in rejected:
        _format_paragraph(doc.add_paragraph(f"{item.get('label') or 'candidate'}: {item.get('reason') or labels['translation_unavailable']}", style="List Bullet"), size=BODY_SIZE)


def _append_watchlist_digest_section(doc: Document, *, artifacts: list[dict[str, str]], language: str, include_title: bool = True) -> None:
    labels = _digest_labels(language)
    if include_title:
        _add_title(doc, labels["title"], subtitle=f"{labels['refreshed_names']}: {len(artifacts)}")
    if not artifacts:
        _format_paragraph(doc.add_paragraph(labels["no_items"]), size=BODY_SIZE)
    for item in artifacts:
        _format_paragraph(doc.add_heading(str(item.get("identifier") or "Name"), level=2), size=HEADING2_SIZE, bold=True)
        _format_paragraph(doc.add_paragraph(f"{labels['verdict']}: {item.get('verdict') or 'watchlist'}", style="List Bullet"), size=BODY_SIZE)
        _format_paragraph(doc.add_paragraph(f"{labels['target_snapshot']}: {item.get('target_snapshot') or 'n/a'}", style="List Bullet"), size=BODY_SIZE)
        path_key = "primary_document_path" if item.get("primary_document_path") else "zh_docx_path"
        _format_paragraph(doc.add_paragraph(f"{labels['document_path']}: {item.get(path_key) or 'n/a'}", style="List Bullet"), size=BODY_SIZE)


def _report_labels(language: str) -> dict[str, str]:
    if language == "en":
        return {
            "memo_title": "Buy-Side Research Memo",
            "ticker": "Ticker",
            "market": "Market",
            "verdict": "Verdict",
            "confidence": "Confidence",
            "exchange": "Exchange",
            "model": "Model",
            "quick_take": "Quick Take",
            "recent_developments": "Recent Developments / Volatility Clues",
            "business_summary": "Business Summary",
            "market_map": "Market Map",
            "china_story": "Angle / Strategic Narrative",
            "sentiment_simulation": "Sentiment Simulation",
            "peer_comparison": "Peer Comparison",
            "committee_takeaways": "Guru Council Notes",
            "scenario_outlook": "Multi-Future Scenario Outlook",
            "bull_case": "Bull Case",
            "bear_case": "Bear Case",
            "catalysts": "Catalysts",
            "risks": "Risks",
            "valuation_view": "Valuation View",
            "target_prices": "Target Prices",
            "horizon": "Horizon",
            "price": "Target Price",
            "thesis": "Thesis",
            "debate_notes": "Red-Team Dissent",
            "evidence": "Evidence Checklist",
            "source": "Source",
            "claim": "Claim",
            "stance": "Stance",
            "next_questions": "Next Questions",
            "no_items": "No fully translated items were available in this run.",
            "short_term": "Short Term",
            "medium_term": "Medium Term",
            "long_term": "Long Term",
            "translation_unavailable": "Translation unavailable in this run.",
            "macro_context": "Macro Context",
            "flow_signal": "Flow Signal",
            "technical_view": "Technical View",
            "factor_exposure": "Factor Exposure",
            "catalyst_calendar": "Catalyst Calendar",
            "factor_value": "Value",
            "factor_momentum": "Momentum",
            "factor_quality": "Quality",
            "factor_size": "Size",
            "factor_volatility": "Volatility",
            "factor_name": "Factor",
            "factor_rating": "Rating",
            "catalyst_event": "Event",
            "catalyst_date": "Date",
            "catalyst_impact": "Impact",
            "catalyst_direction": "Direction",
        }
    return {
        "memo_title": "研究备忘录",
        "ticker": "标的",
        "market": "市场",
        "verdict": "结论",
        "confidence": "信心",
        "exchange": "交易所",
        "model": "模型",
        "quick_take": "快速判断",
        "recent_developments": "最新实效信息与波动线索",
        "business_summary": "业务概览",
        "market_map": "市场与行业图谱",
        "china_story": "研究角度 / 叙事主线",
        "sentiment_simulation": "舆情与叙事模拟",
        "peer_comparison": "横向对比与估值锚",
        "committee_takeaways": "股神议会纪要",
        "scenario_outlook": "多未来场景",
        "bull_case": "多头逻辑",
        "bear_case": "空头逻辑",
        "catalysts": "催化剂",
        "risks": "主要风险",
        "valuation_view": "估值视角",
        "target_prices": "目标价与时间框架",
        "horizon": "时间",
        "price": "目标价",
        "thesis": "依据",
        "debate_notes": "投委会红队质询",
        "evidence": "证据清单",
        "source": "来源",
        "claim": "核心论点",
        "stance": "立场",
        "next_questions": "下一步验证问题",
        "no_items": "暂无充分结论，需继续补证。",
        "short_term": "短期",
        "medium_term": "中期",
        "long_term": "长期",
        "translation_unavailable": "当前证据仍不足，需要继续核实。",
        "macro_context": "宏观环境",
        "flow_signal": "资金流向信号",
        "technical_view": "技术面视角",
        "factor_exposure": "因子暴露",
        "catalyst_calendar": "催化剂日历",
        "factor_value": "价值",
        "factor_momentum": "动量",
        "factor_quality": "质量",
        "factor_size": "规模",
        "factor_volatility": "波动率",
        "factor_name": "因子",
        "factor_rating": "评级",
        "catalyst_event": "事件",
        "catalyst_date": "日期",
        "catalyst_impact": "影响",
        "catalyst_direction": "方向",
    }


def _screening_labels(language: str) -> dict[str, str]:
    if language == "en":
        return {
            "title": "Screening Brief",
            "market": "Market",
            "second_screen_pool": "Second-Screen Pool",
            "final_recommendations": "Final Recommendations",
            "recommendation_rank": "Recommendation Rank",
            "screen_score": "Screen Score",
            "research_verdict": "Research Verdict",
            "confidence": "Confidence",
            "why_now": "Why Now",
            "why_not_now": "Why Not Now",
            "quick_take": "Quick Take",
            "vertical_summary": "Vertical Diligence",
            "horizontal_summary": "Horizontal Diligence",
            "short_term_target": "Short-Term Target",
            "bull_bear_focus": "Bull/Bear Focus",
            "document_path": "Report Document",
            "downgraded_names": "Downgraded or Excluded Names",
            "translation_unavailable": "English translation was unavailable in this run.",
            "no_items": "No items were available in this bucket.",
        }
    return {
        "title": "筛股报告",
        "market": "市场",
        "second_screen_pool": "二筛候选池",
        "final_recommendations": "最终推荐",
        "recommendation_rank": "推荐等级",
        "screen_score": "筛选分数",
        "research_verdict": "研究结论",
        "confidence": "信心",
        "why_now": "为什么现在值得看",
        "why_not_now": "为什么现在还不能下重注",
        "quick_take": "快速判断",
        "vertical_summary": "纵向尽调",
        "horizontal_summary": "横向尽调",
        "short_term_target": "短期目标价",
        "bull_bear_focus": "多空焦点",
        "document_path": "报告文档",
        "downgraded_names": "降级或淘汰名单",
        "translation_unavailable": "当前轮次还没有更清晰的补充说明。",
        "no_items": "当前桶里没有结果。",
    }


def _digest_labels(language: str) -> dict[str, str]:
    if language == "en":
        return {
            "title": "Watchlist Digest",
            "refreshed_names": "Refreshed Names",
            "verdict": "Verdict",
            "target_snapshot": "Target Snapshot",
            "document_path": "Report Document",
            "no_items": "No watchlist names were due in this cycle.",
        }
    return {
        "title": "跟踪池更新摘要",
        "refreshed_names": "本轮刷新标的数",
        "verdict": "结论",
        "target_snapshot": "目标价快照",
        "document_path": "报告文档",
        "no_items": "本轮没有到期需要刷新的标的。",
    }
